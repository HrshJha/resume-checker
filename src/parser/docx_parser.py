"""
DOCX resume parser using python-docx.

Extracts text from paragraphs, tables, and nested lists.
Preserves structure for section detection.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from src.utils.logger import get_logger

logger = get_logger("docx_parser")


@dataclass
class DOCXParseResult:
    """Result of parsing a DOCX file."""

    full_text: str
    paragraphs: list[dict] = field(default_factory=list)
    tables: list[list[list[str]]] = field(default_factory=list)
    warnings: list[str] = field(default_factory=list)


def parse_docx(file_path: str | Path) -> DOCXParseResult:
    """
    Parse a DOCX file and extract structured text.

    Args:
        file_path: Path to the DOCX file.

    Returns:
        DOCXParseResult with paragraphs, tables, and full text.
    """
    from docx import Document

    file_path = Path(file_path)
    warnings: list[str] = []

    try:
        doc = Document(str(file_path))
    except Exception as e:
        logger.error(f"Failed to open DOCX {file_path}: {e}")
        return DOCXParseResult(
            full_text="",
            warnings=[f"Failed to open DOCX: {e}"],
        )

    # Extract paragraphs with style information
    paragraphs: list[dict] = []
    text_parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""
        is_heading = style_name.lower().startswith("heading")
        level = 0
        if is_heading:
            try:
                level = int(style_name.split()[-1])
            except (ValueError, IndexError):
                level = 1

        # Detect list items
        is_list_item = False
        if para.paragraph_format and para.paragraph_format.first_line_indent:
            is_list_item = True
        if text.startswith(("•", "-", "–", "▪", "○", "●")):
            is_list_item = True
            text = text.lstrip("•-–▪○● ").strip()
        if text and text[0].isdigit() and "." in text[:4]:
            is_list_item = True

        # Check for bold text (potential section header)
        is_bold = False
        if para.runs:
            is_bold = all(run.bold for run in para.runs if run.text.strip())

        paragraphs.append(
            {
                "text": text,
                "style": style_name,
                "is_heading": is_heading,
                "heading_level": level,
                "is_bold": is_bold,
                "is_list_item": is_list_item,
            }
        )
        text_parts.append(text)

    # Extract tables
    tables: list[list[list[str]]] = []
    for table in doc.tables:
        table_data: list[list[str]] = []
        for row in table.rows:
            row_data: list[str] = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                row_data.append(cell_text)
            if any(cell for cell in row_data):
                table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    full_text = "\n".join(text_parts)

    if not full_text.strip():
        warnings.append("No text content extracted from DOCX")

    logger.debug(
        f"Parsed DOCX {file_path.name}: {len(paragraphs)} paragraphs, {len(tables)} tables"
    )

    return DOCXParseResult(
        full_text=full_text,
        paragraphs=paragraphs,
        tables=tables,
        warnings=warnings,
    )
